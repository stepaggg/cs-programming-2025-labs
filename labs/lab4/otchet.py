# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import subprocess

def open_file_in_word(filename):
    """Автоматическое открытие файла в Word"""
    try:
        if os.name == 'nt':  # Windows
            os.startfile(filename)
        else:  # Mac или Linux
            subprocess.call(['open', filename])
        print(f"✓ Открываю файл: {filename}")
    except Exception as e:
        print(f"✗ Не удалось открыть файл автоматически: {e}")

def create_complete_report():
    """Создание ПОЛНОГО отчета с титульным листом и всей работой"""
    
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    
    # Настройка страницы
    section = doc.sections[0]
    section.page_height = Inches(11.7)
    section.page_width = Inches(8.3)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(1.2)
    section.bottom_margin = Inches(1.2)
    
    # ===== ТОЧНЫЙ ТИТУЛЬНЫЙ ЛИСТ =====
    print("Создаю титульный лист...")
    
    # Верхние отступы
    for _ in range(10):
        doc.add_paragraph()
    
    # МИНОБРНАУКИ РОССИИ (выровнено по левому краю как в образце)
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = p1.add_run("МИНОБРНАУКИ РОССИИ")
    run1.font.size = Pt(14)
    run1.bold = True
    run1.font.name = 'Times New Roman'
    
    # Федеральное государственное... (выровнено по левому краю)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = p2.add_run("Федеральное государственное бюджетное образовательное учреждение")
    run2.font.size = Pt(14)
    run2.font.name = 'Times New Roman'
    
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run3 = p3.add_run("высшего образования")
    run3.font.size = Pt(14)
    run3.font.name = 'Times New Roman'
    
    # ВГУ в кавычках (выровнено по левому краю)
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run4 = p4.add_run("«ВЛАДИВОСТОКСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ»")
    run4.font.size = Pt(14)
    run4.bold = True
    run4.font.name = 'Times New Roman'
    
    # ФГБОУ ВО ВВГУ (выровнено по левому краю)
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run5 = p5.add_run("(ФГБОУ ВО «ВВГУ»)")
    run5.font.size = Pt(14)
    run5.font.name = 'Times New Roman'
    
    # Отступ
    for _ in range(2):
        doc.add_paragraph()
    
    # Институт и кафедра (выровнено по левому краю)
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run6 = p6.add_run("ИНСТИТУТ ИНФОРМАЦИОННЫХ ТЕХНОЛОГИЙ И АНАЛИЗА ДАННЫХ")
    run6.font.size = Pt(14)
    run6.bold = True
    run6.font.name = 'Times New Roman'
    
    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run7 = p7.add_run("КАФЕДРА ИНФОРМАЦИОННЫХ ТЕХНОЛОГИЙ И СИСТЕМ")
    run7.font.size = Pt(14)
    run7.bold = True
    run7.font.name = 'Times New Roman'
    
    # Разделительная линия
    for _ in range(3):
        doc.add_paragraph()
    
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("―" * 50)
    run_line.font.size = Pt(14)
    run_line.font.name = 'Times New Roman'
    
    # Отступ после линии
    for _ in range(3):
        doc.add_paragraph()
    
    # ОТЧЕТ (по центру)
    p8 = doc.add_paragraph()
    p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run8 = p8.add_run("ОТЧЕТ")
    run8.font.size = Pt(16)
    run8.bold = True
    run8.font.name = 'Times New Roman'
    
    # ПО ЛАБОРАТОРНОЙ РАБОТЕ (по центру)
    p9 = doc.add_paragraph()
    p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run9 = p9.add_run("ПО ЛАБОРАТОРНОЙ РАБОТЕ №1")
    run9.font.size = Pt(14)
    run9.bold = True
    run9.font.name = 'Times New Roman'
    
    # по дисциплине (по центру)
    p10 = doc.add_paragraph()
    p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run10 = p10.add_run("по дисциплине")
    run10.font.size = Pt(14)
    run10.font.name = 'Times New Roman'
    
    # Название дисциплины (по центру)
    p11 = doc.add_paragraph()
    p11.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run11 = p11.add_run("«Информатика и программирование»")
    run11.font.size = Pt(14)
    run11.bold = True
    run11.font.name = 'Times New Roman'
    
    # Большой отступ перед данными студента
    for _ in range(8):
        doc.add_paragraph()
    
    # Студент (по правому краю)
    p12 = doc.add_paragraph()
    p12.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run12 = p12.add_run("Студент")
    run12.font.size = Pt(14)
    run12.font.name = 'Times New Roman'
    
    # Группа БИН-25-2 (по правому краю)
    p13 = doc.add_paragraph()
    p13.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run13 = p13.add_run("гр. БИН-25-2")
    run13.font.size = Pt(14)
    run13.font.name = 'Times New Roman'
    
    # Ассистент (по правому краю)
    p14 = doc.add_paragraph()
    p14.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run14 = p14.add_run("Ассистент")
    run14.font.size = Pt(14)
    run14.font.name = 'Times New Roman'
    
    # преподавателя (по правому краю)
    p15 = doc.add_paragraph()
    p15.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run15 = p15.add_run("преподавателя")
    run15.font.size = Pt(14)
    run15.font.name = 'Times New Roman'
    
    # Отступ для подписей
    for _ in range(4):
        doc.add_paragraph()
    
    # Линия для подписи студента (по правому краю)
    p16 = doc.add_paragraph()
    p16.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run16 = p16.add_run("_________________________")
    run16.font.size = Pt(14)
    run16.font.name = 'Times New Roman'
    
    # ФИО студента С.А. Головцов (по правому краю)
    p17 = doc.add_paragraph()
    p17.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run17 = p17.add_run("С.А. Головцов")
    run17.font.size = Pt(14)
    run17.font.name = 'Times New Roman'
    
    # Линия для подписи преподавателя (по правому краю)
    p18 = doc.add_paragraph()
    p18.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run18 = p18.add_run("_________________________")
    run18.font.size = Pt(14)
    run18.font.name = 'Times New Roman'
    
    # ФИО преподавателя (по правому краю)
    p19 = doc.add_paragraph()
    p19.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run19 = p19.add_run("М.В. Водяницкий")
    run19.font.size = Pt(14)
    run19.font.name = 'Times New Roman'
    
    # Отступ перед городом и годом
    for _ in range(8):
        doc.add_paragraph()
    
    # Разделительная линия
    p_line2 = doc.add_paragraph()
    p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line2 = p_line2.add_run("―" * 50)
    run_line2.font.size = Pt(14)
    run_line2.font.name = 'Times New Roman'
    
    # Отступ после линии
    for _ in range(2):
        doc.add_paragraph()
    
    # Город и год (жирный)
    p20 = doc.add_paragraph()
    p20.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run20 = p20.add_run("Владивосток 2025")
    run20.font.size = Pt(14)
    run20.bold = True
    run20.font.name = 'Times New Roman'
    
    # ===== СОДЕРЖАНИЕ =====
    doc.add_page_break()
    print("Создаю содержание...")
    
    p_content = doc.add_paragraph()
    p_content.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_content = p_content.add_run("СОДЕРЖАНИЕ")
    run_content.font.size = Pt(16)
    run_content.bold = True
    run_content.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    content_items = [
        "1. Введение........................................................3",
        "2. Задание 1: Управление кондиционером............................4",
        "3. Задание 2: Определение времени года............................5", 
        "4. Задание 3: Расчет возраста собаки..............................6",
        "5. Задание 4: Проверка делимости на 6.............................7",
        "6. Задание 5: Проверка надежности пароля..........................8",
        "7. Задание 6: Определение високосного года........................9",
        "8. Задание 7: Поиск наименьшего числа............................10",
        "9. Задание 8: Расчет скидки в магазине...........................11",
        "10. Задание 9: Определение времени суток.........................12",
        "11. Задание 10: Проверка простого числа..........................13",
        "12. Заключение....................................................14",
        "13. Список использованных источников.............................15"
    ]
    
    for item in content_items:
        p = doc.add_paragraph(item)
        p.style = doc.styles['Normal']
    
    # ===== ВВЕДЕНИЕ =====
    doc.add_page_break()
    print("Создаю раздел 'Введение'...")
    
    p_intro = doc.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_intro = p_intro.add_run("1. ВВЕДЕНИЕ")
    run_intro.font.size = Pt(16)
    run_intro.bold = True
    run_intro.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    intro_text = [
        "Цель лабораторной работы: освоение базовых конструкций языка Python",
        "и разработка практических задач с использованием условных операторов.",
        "",
        "Задачи работы:",
        "1. Изучение синтаксиса Python",
        "2. Разработка 10 практических задач", 
        "3. Освоение работы с условными операторами if-else",
        "4. Разработка программ с пользовательским вводом",
        "5. Анализ эффективности разработанных алгоритмов",
        "",
        "Лабораторная работа выполнена в среде программирования Python 3.x",
        "с использованием стандартных библиотек языка."
    ]
    
    for text in intro_text:
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing = 1.5
    
    # ===== ЗАДАНИЯ =====
    tasks = [
        {
            "title": "2. ЗАДАНИЕ 1: УПРАВЛЕНИЕ КОНДИЦИОНЕРОМ",
            "code": '''def task1():
    temp = float(input("Введите температуру: "))
    print("Кондиционер выключен" if temp >= 20 else "Кондиционер включен")''',
            "explanation": "Программа управляет кондиционером на основе температуры. При температуре 20°C и выше кондиционер выключается, при меньшей температуре - включается."
        },
        {
            "title": "3. ЗАДАНИЕ 2: ОПРЕДЕЛЕНИЕ ВРЕМЕНИ ГОДА", 
            "code": '''def task2():
    month = int(input("Введите номер месяца: "))
    if month in [12, 1, 2]:
        print("Это зима")
    elif month in [3, 4, 5]:
        print("Это весна")
    elif month in [6, 7, 8]:
        print("Это лето")
    else:
        print("Это осень")''',
            "explanation": "Программа определяет время года по номеру месяца. Зима: 12,1,2; Весна: 3,4,5; Лето: 6,7,8; Осень: 9,10,11."
        },
        {
            "title": "4. ЗАДАНИЕ 3: РАСЧЕТ ВОЗРАСТА СОБАКИ",
            "code": '''def task3():
    age = float(input("Введите возраст собаки: "))
    if age <= 2:
        result = age * 10.5
    else:
        result = 21 + (age - 2) * 4
    print(f"Возраст в человеческих годах: {result}")''',
            "explanation": "Перевод возраста собаки в человеческие годы. Первые 2 года: каждый год = 10.5 человеческих лет, последующие: каждый год = 4 человеческих года."
        },
        {
            "title": "5. ЗАДАНИЕ 4: ПРОВЕРКА ДЕЛИМОСТИ НА 6",
            "code": '''def task4():
    num = input("Введите число: ")
    last_digit = int(num[-1])
    digit_sum = sum(int(d) for d in num)
    
    if last_digit % 2 == 0 and digit_sum % 3 == 0:
        print("Делится на 6")
    else:
        print("Не делится на 6")''',
            "explanation": "Проверка делимости на 6. Число делится на 6 если: оно четное (последняя цифра делится на 2) И сумма цифр делится на 3."
        },
        {
            "title": "6. ЗАДАНИЕ 5: ПРОВЕРКА НАДЕЖНОСТИ ПАРОЛЯ",
            "code": '''def task5():
    pwd = input("Введите пароль: ")
    errors = []
    
    if len(pwd) < 8:
        errors.append("длина менее 8 символов")
    if not any(c.isupper() for c in pwd):
        errors.append("нет заглавных букв")
    if not any(c.islower() for c in pwd):
        errors.append("нет строчных букв")
    if not any(c.isdigit() for c in pwd):
        errors.append("нет цифр")
    if pwd.isalnum():
        errors.append("нет спецсимволов")
    
    if errors:
        print("Пароль ненадежный:", ", ".join(errors))
    else:
        print("Пароль надежный")''',
            "explanation": "Проверка пароля по критериям: длина ≥8, есть заглавные и строчные буквы, есть цифры, есть спецсимволы."
        }
    ]
    
    # Добавляем первые 5 заданий
    for i, task in enumerate(tasks):
        doc.add_page_break()
        print(f"Создаю {task['title']}...")
        
        # Заголовок задания
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run(task['title'])
        run_title.font.size = Pt(14)
        run_title.bold = True
        run_title.font.name = 'Times New Roman'
        
        doc.add_paragraph()
        
        # Код задания
        p_code_title = doc.add_paragraph("Код программы:")
        p_code_title.style = doc.styles['Normal']
        
        p_code = doc.add_paragraph(task['code'])
        p_code.style = doc.styles['Normal']
        
        doc.add_paragraph()
        
        # Объяснение
        p_expl = doc.add_paragraph("Объяснение:")
        p_expl.style = doc.styles['Normal']
        
        p_expl_text = doc.add_paragraph(task['explanation'])
        p_expl_text.style = doc.styles['Normal']
        p_expl_text.paragraph_format.line_spacing = 1.5
    
    # ===== ОСТАЛЬНЫЕ ЗАДАНИЯ =====
    more_tasks = [
        {
            "number": "7",
            "title": "ЗАДАНИЕ 6: ОПРЕДЕЛЕНИЕ ВИСОКОСНОГО ГОДА",
            "code": '''def task6():
    year = int(input("Введите год: "))
    if (year % 4 == 0 and year % 100 != 0) or year % 400 == 0:
        print(f"{year} - високосный")
    else:
        print(f"{year} - невисокосный")''',
            "explanation": "Определение високосного года. Год високосный если: делится на 4 И не делится на 100, ИЛИ делится на 400."
        },
        {
            "number": "8", 
            "title": "ЗАДАНИЕ 7: ПОИСК НАИМЕНЬШЕГО ЧИСЛА",
            "code": '''def task7():
    a, b, c = map(float, input("Введите три числа: ").split())
    min_num = a
    if b < min_num:
        min_num = b
    if c < min_num:
        min_num = c
    print(f"Наименьшее: {min_num}")''',
            "explanation": "Поиск наименьшего из трех чисел. Алгоритм последовательно сравнивает числа и обновляет минимальное значение."
        },
        {
            "number": "9",
            "title": "ЗАДАНИЕ 8: РАСЧЕТ СКИДКИ В МАГАЗИНЕ", 
            "code": '''def task8():
    amount = float(input("Введите сумму покупки: "))
    if amount < 1000:
        discount = 0
    elif amount <= 5000:
        discount = 5
    elif amount <= 10000:
        discount = 10
    else:
        discount = 15
    
    total = amount * (1 - discount / 100)
    print(f"Скидка: {discount}%")
    print(f"К оплате: {total}")''',
            "explanation": "Расчет скидки в зависимости от суммы покупки: до 1000р - 0%, 1000-5000р - 5%, 5000-10000р - 10%, свыше 10000р - 15%."
        },
        {
            "number": "10",
            "title": "ЗАДАНИЕ 9: ОПРЕДЕЛЕНИЕ ВРЕМЕНИ СУТОК",
            "code": '''def task9():
    hour = int(input("Введите час (0-23): "))
    if 0 <= hour <= 5:
        print("Ночь")
    elif 6 <= hour <= 11:
        print("Утро")
    elif 12 <= hour <= 17:
        print("День")
    else:
        print("Вечер")''',
            "explanation": "Определение времени суток по часу: 0-5 - ночь, 6-11 - утро, 12-17 - день, 18-23 - вечер."
        },
        {
            "number": "11",
            "title": "ЗАДАНИЕ 10: ПРОВЕРКА ПРОСТОГО ЧИСЛА", 
            "code": '''def task10():
    num = int(input("Введите число: "))
    if num < 2:
        print("Не простое")
        return
    
    for i in range(2, int(num**0.5) + 1):
        if num % i == 0:
            print(f"{num} - составное")
            return
    print(f"{num} - простое")''',
            "explanation": "Проверка числа на простоту. Число простое если делится только на 1 и на себя. Проверка делителей от 2 до √n."
        }
    ]
    
    # Добавляем остальные задания
    for task in more_tasks:
        doc.add_page_break()
        print(f"Создаю задание {task['number']}...")
        
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run(f"{task['number']}. {task['title']}")
        run_title.font.size = Pt(14)
        run_title.bold = True
        run_title.font.name = 'Times New Roman'
        
        doc.add_paragraph()
        
        p_code_title = doc.add_paragraph("Код программы:")
        p_code = doc.add_paragraph(task['code'])
        
        doc.add_paragraph()
        
        p_expl_title = doc.add_paragraph("Объяснение:")
        p_expl = doc.add_paragraph(task['explanation'])
        p_expl.paragraph_format.line_spacing = 1.5
    
    # ===== ЗАКЛЮЧЕНИЕ =====
    doc.add_page_break()
    print("Создаю заключение...")
    
    p_conc = doc.add_paragraph()
    p_conc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_conc = p_conc.add_run("12. ЗАКЛЮЧЕНИЕ")
    run_conc.font.size = Pt(16)
    run_conc.bold = True
    run_conc.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    conclusion_text = [
        "В ходе выполнения лабораторной работы были успешно решены 10 практических",
        "задач на языке Python. Были освоены следующие навыки:",
        "",
        "1. Работа с условными операторами if-elif-else",
        "2. Обработка пользовательского ввода",
        "3. Использование тернарных операторов", 
        "4. Работа со строками и преобразование типов данных",
        "5. Реализация математических алгоритмов",
        "6. Проверка различных условий и критериев",
        "",
        "Все задачи были решены оптимально с точки зрения алгоритмической сложности.",
        "Разработанные программы демонстрируют практическое применение условных",
        "операторов в реальных сценариях.",
        "",
        "Результаты работы показывают уверенное владение базовыми конструкциями",
        "языка Python и умение применять их для решения практических задач."
    ]
    
    for text in conclusion_text:
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing = 1.5
    
    # ===== СПИСОК ИСТОЧНИКОВ =====
    doc.add_page_break()
    print("Создаю список источников...")
    
    p_bibl = doc.add_paragraph()
    p_bibl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_bibl = p_bibl.add_run("13. СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    run_bibl.font.size = Pt(16)
    run_bibl.bold = True
    run_bibl.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    sources = [
        "1. ГОСТ 7.32-2001 'Отчет о научно-исследовательской работе'",
        "2. ГОСТ 7.0.5-2008 'Библиографическая запись'", 
        "3. Официальная документация Python: https://docs.python.org",
        "4. Владивостокский государственный университет. Стандарты оформления работ",
        "5. Лутц М. 'Изучаем Python'. - СПб.: Символ-Плюс, 2022",
        "6. Python Software Foundation. Python 3.11 Documentation"
    ]
    
    for source in sources:
        p = doc.add_paragraph(source)
        p.paragraph_format.line_spacing = 1.5
    
    # Сохраняем полный документ
    filename = "ПОЛНЫЙ_ОТЧЕТ_ЛАБОРАТОРНАЯ_1.docx"
    doc.save(filename)
    
    print(f"✓ Полный отчет создан: {filename}")
    print(f"✓ Объем: около 15-20 страниц")
    return filename

def install_requirements():
    """Инструкция по установке библиотек"""
    print("\nДля установки библиотеки выполните:")
    print("pip install python-docx")

if __name__ == "__main__":
    try:
        print("=" * 60)
        print("СОЗДАНИЕ ПОЛНОГО ОТЧЕТА С ТИТУЛЬНЫМ ЛИСТОМ")
        print("=" * 60)
        
        # Создаем полный отчет
        filename = create_complete_report()
        
        print("\n" + "=" * 60)
        print("✓ ОТЧЕТ УСПЕШНО СОЗДАН!")
        print("=" * 60)
        print(f"Файл: {filename}")
        print(f"Размер: {os.path.getsize(filename)} байт")
        print(f"Путь: {os.path.abspath(filename)}")
        print(f"Студент: С.А. Головцов")
        print(f"Группа: БИН-25-2")
        
        # Открываем файл
        print("\nОткрываю файл в Word...")
        open_file_in_word(filename)
        
        print("\n" + "=" * 60)
        print("СОДЕРЖАНИЕ ОТЧЕТА:")
        print("- Точный титульный лист (как в образце)")
        print("- Содержание") 
        print("- Введение")
        print("- 10 заданий с кодом и объяснениями")
        print("- Заключение")
        print("- Список источников")
        print("=" * 60)
        
    except ImportError:
        print("✗ Библиотека python-docx не установлена!")
        install_requirements()
    except Exception as e:
        print(f"✗ Ошибка при создании отчета: {e}")
        install_requirements()